#!usr/bin/env python  
#-*- coding:utf-8 _*-  

"""
@Author: SMnRa
@Email: smnra@163.com
@Project: docx
@File: readDocx.py
@Time: 2019/04/27 9:37

功能描述:
        教案格式转化



"""
import sys,os

from docx import Document
from docx import shared
from docx.shared import RGBColor

from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.style import WD_STYLE_TYPE
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_TABLE_DIRECTION





def doc2doc(sourceFile,tagDir):

    # 打开文档
    tagDocument = Document()

    # 新增自定义样式
    styles = tagDocument.styles

    paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH][0]
    character_styles = [s for s in styles if s.type == WD_STYLE_TYPE.CHARACTER][0]
    table_styles = [s for s in styles if s.type == WD_STYLE_TYPE.TABLE][0]

    # 设置与上一段间隔 Pt（5）
    paragraph_styles.paragraph_format.space_after = shared.Pt(5)
    # 设置与下一段间隔 Pt（10）
    paragraph_styles.paragraph_format.space_before = shared.Pt(10)
    # 行间距
    paragraph_styles.paragraph_format.line_spacing = shared.Pt(18)


    paragraph_styles.font.name = u'宋体'
    paragraph_styles._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    paragraph_styles.font.size = shared.Pt(15)
    paragraph_styles.font.color.rgb = RGBColor(0,0,0)
    paragraph_styles.font.underline = False

    character_styles.font.name = u'宋体'
    character_styles.font.size = shared.Pt(15)
    character_styles.font.underline = True

    # table_styles.






    # 加入不同等级的标题
    titleParagraph = tagDocument.add_paragraph(style=paragraph_styles)
    # 段落对齐方式
    paragraph_format = titleParagraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    titleRun = titleParagraph.add_run(u'西安市教师专用教案')
    titleRun.font.name = u'黑体'
    titleRun._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    titleRun.font.size = shared.Pt(22)   # 字号
    titleRun.font.bold  = True    # 加粗
    titleRun.font.color.rgb = RGBColor(0,0,0)





    # 添加段落 (块等级)
    paragraph = tagDocument.add_paragraph()
    # 段落对齐方式
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



    #内联等级设置字号
    run = paragraph.add_run(u' 二 ',style=character_styles)
    run2 = paragraph.add_run(u'年级  2018年－2019学年度  第二学期  第1周  教师',style=character_styles)
    run3 = paragraph.add_run(u' 石林  ',style=character_styles)

    run.font.size = shared.Pt(14)
    run.font.underline = True
    run2.font.size = shared.Pt(14)
    run3.font.size = shared.Pt(14)
    run3.font.underline = True


    def createTable(course,tagDir):
        tagName = u"教案_" + course[0][1] + ".docx"

        # 创建目标格式的表格
        table = tagDocument.add_table(19,6, style = "Table Grid")

        # 当前正在操作的行号
        rowindex = 0
        # #可以使表格自动适应窗口大小。
        # table.autofit=True

        # 设置单元格宽度
        # table.cell(row,col).width=shared.Pt(82)

        # 设置行高度
        # table.rows[0].height=shared.Pt(82)

        # 设置列宽
        # table.columns[0].width = shared.Pt(82)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        table.style.font.size=shared.Pt(15)
        table.style.font.name='宋体'
        table.style.font.color.rgb=RGBColor(0,0,123)


        # 将cell(0,0)到cell(2,2)之间的所有cell合并成一个cell
        # table.cell(0,0).merge(table.cell(2,2))
        # 这里需要注意的是，虽然每个cell都合并了，但其实它还是存在的。比如合并了(0,0)和(0,1)两个cell，那么这个合并的cell其实就是(0,0;0,1)



        # 设置列宽
        table.columns[0].width = shared.Pt(60)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER


        #  列表头填充
        table.cell(rowindex,0).text = "班级"
        table.cell(rowindex,1).text = "二二班"
        table.cell(rowindex,2).text = "科目"
        table.cell(rowindex,3).text = "书法"
        table.cell(rowindex,4).text = "教学时数"
        table.cell(rowindex,5).text = "  1 课时"
        rowindex+=1

        # 课题 单元格合并
        cell_keti = table.cell(rowindex,0)
        cell_keti.text = "课题"

        # 课题 的内容 单元格合并
        cell_keti_neirong = table.cell(1,1).merge(table.cell(1,5))
        cell_keti_neirong.text = course[0][1]
        rowindex += 1


        # 教学目的要求 的内容 单元格合并
        ketineirongs = course[1][1].split("\n")
        for i, ketineirong in enumerate(ketineirongs):
            cell_ketineirong = table.cell(rowindex+i,1).merge(table.cell(rowindex+i,5))
            cell_ketineirong.text = ketineirong.strip()
        else:
            print(rowindex,rowindex+i)
            # 教学目的要求 单元格合并
            cell_mudi = table.cell(rowindex, 0).merge(table.cell(rowindex+i, 0))
            cell_mudi.text = """教学\n目的\n要求"""
            rowindex += i + 1


        # 重点难点 的内容 单元格合并
        zhongdians = course[2][1].split("，")
        for i, zhongdian in enumerate(zhongdians):
            cell_zhongdian_neirong = table.cell(rowindex+i,1).merge(table.cell(rowindex+i,5))
            cell_zhongdian_neirong.text = zhongdian.strip()
        else:
            print(rowindex, rowindex + i)
            # 重点难点 单元格合并
            cell_zhongdians = table.cell(rowindex, 0).merge(table.cell(rowindex+i, 0))
            cell_zhongdians.text = """重点\n难点"""
            rowindex += i +1


        # 教学方法 单元格
        cell_fangfa = table.cell(rowindex,0)
        cell_fangfa.text = "教学方法"

        # 教学方法  的内容 单元格合并
        cell_fangfa_neirong = table.cell(rowindex,1).merge(table.cell(rowindex,5))
        cell_fangfa_neirong.text = "讲授法   讨论法"
        rowindex += 1


        # 教具 单元格
        cell_jiaoju = table.cell(rowindex,0)
        cell_jiaoju.text = "教具"

        # 教具  的内容 单元格合并
        cell_jiaoju_neirong = table.cell(rowindex,1).merge(table.cell(rowindex,5))
        cell_jiaoju_neirong.text = "教学课件"
        rowindex += 1


        # 复习检查 单元格
        cell_fuxi = table.cell(rowindex,0)
        cell_fuxi.text = """复习\n检查"""

        # 复习检查 的内容 单元格合并
        cell_fuxi_neirong = table.cell(rowindex,1).merge(table.cell(rowindex,5))
        cell_fuxi_neirong.text = ""
        rowindex += 1


        # 板书设计 单元格
        cell_sheji = table.cell(rowindex,0)
        cell_sheji.text = """板\n书\n设\n计"""

        # 板书设计 的内容 单元格合并
        cell_sheji_neirong = table.cell(rowindex,1).merge(table.cell(rowindex,5))
        cell_sheji_neirong.text = ""
        rowindex += 1


        # 课堂练习 课后作业 实践活动  单元格
        cell_lianxi = table.cell(rowindex,0)
        cell_lianxi.text = """课堂练习\n课后作业\n实践活动"""

        # "课堂练习 课后作业 实践活动" 的内容 单元格合并
        cell_lianxi_neirong = table.cell(rowindex,1).merge(table.cell(rowindex,5))
        cell_lianxi_neirong.text = ""
        rowindex += 1


        # 课后记  单元格
        cell_kehouji = table.cell(rowindex,0)
        cell_kehouji.text = """课后记"""

        # "课堂练习 课后作业 实践活动" 的内容 单元格合并
        cell_kehouji_neirong = table.cell(rowindex,1).merge(table.cell(rowindex,5))
        cell_kehouji_neirong.text = ""
        rowindex += 1


        # 三个空行 单元格
        cell_konghang_1 = table.cell(rowindex,0).merge(table.cell(rowindex,5))
        cell_konghang_2 = table.cell(rowindex+1,0).merge(table.cell(rowindex+1,5))
        cell_konghang_3 = table.cell(rowindex+2, 0).merge(table.cell(rowindex+2, 5))
        rowindex += 3


        # "教 学 程 序 及 内 容"  单元格
        cell_neirong_title = table.cell(rowindex,0).merge(table.cell(rowindex,3))
        cell_neirong_title.text = """教 学 程 序 及 内 容"""

        # "学生活动设计"  单元格合并
        cell_huodong_title = table.cell(rowindex,4).merge(table.cell(rowindex,5))
        cell_huodong_title.text = """学生活动设计"""
        rowindex += 1


        # "教 学 程 序 及 内 容"  的内容 单元格合并
        cell_neirong_title = table.cell(rowindex,0).merge(table.cell(rowindex,3))
        cell_neirong_title.text = "\n".join(course[4])

        # "学生活动设计"  的内容 单元格合并
        cell_huodong_title = table.cell(rowindex,4).merge(table.cell(rowindex,5))
        cell_huodong_title.text = ""
        rowindex += 1
        table.columns[0].width = shared.Pt(60)


        #保存文档
        tagDocument.save(tagDir + r'/' + tagName)
        print(tagName, "is OK.")





    #打开文档
    # print(len(sys.argv))
    # print(sys.argv[0])
    # print(sys.argv[1])
    # print(sys.argv[2])
    if len(sys.argv) == 3:
        sourceFileName = sys.argv[1]
        tagPath = sys.argv[2]
    else:
        # sourceFileName = u'./doc/二年级下册书法教案.docx'
        sourceFileName = sourceFile

        # tagPath = r"./doc"
        tagPath = tagDir


    sourceDocument = Document(sourceFileName)
    #  注意此处 打开文档的路径字符串 应为 unicode 字符串 u'.\\doc\\二年级下册书法教案.docx'

    #读取每段资料
    l = [ paragraph.text.encode('gb2312') for paragraph in sourceDocument.paragraphs]

    #输出文档中的文本
    for i in l:
        print(i.decode('gb2312'))


    # 读取表格:
    tables = sourceDocument.tables

    # 课程列表
    courseList = []

    for table in tables:
        rowCount = table.rows.__len__()      # 表格的总行数
        colCount = table.columns.__len__()    #表格的总列数
        rows = []
        row = []
        for r in range(0,rowCount):
            row = []

            # print("row: "+str(r+1))
            for c in range(0,colCount):
                cellText = table.cell(r,c).text
                # print(cellText)
                if cellText not in row:
                    row.append(cellText)
            if row not in rows :
                rows.append(row)
            # cell(i,0)表示第(i+1)行第1列数据，以此类推
            # print(result)
            # print(result_1)

        print(rows)
        courseList.append(rows)




    for course in courseList:
        createTable(course,tagDir)

    print("Over!")